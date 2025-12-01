import streamlit as st
import pandas as pd
import numpy as np
import scipy.stats as stats
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.stats.power import TTestIndPower
from lifelines import KaplanMeierFitter, CoxPHFitter
from sklearn.metrics import roc_curve, auc, confusion_matrix, classification_report
from sklearn.impute import SimpleImputer
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
import matplotlib.pyplot as plt
import seaborn as sns
import base64
import io
from datetime import datetime
import time
import random
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# CONFIGURATION & GLOBAL STYLES
# ==========================================
st.set_page_config(
    page_title="MedStat Pro | Powerhouse Research Suite",
    page_icon="⚕️",
    layout="wide",
    initial_sidebar_state="expanded"
)

APP_NAME = "MedStat Pro"
VERSION = "3.0 (Powerhouse)"

# Custom CSS for Professional UI
st.markdown("""
<style>
    .main { background-color: #f4f6f9; }
    h1, h2, h3 { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #2c3e50; }
    .stButton>button { width: 100%; border-radius: 6px; font-weight: 600; height: 3em; transition: all 0.3s ease; }
    .stButton>button:hover { border-color: #007bff; color: white; background-color: #007bff; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
    .report-card { background: white; padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 20px; }
    .metric-container { background: #e3f2fd; padding: 15px; border-radius: 8px; text-align: center; border: 1px solid #90caf9; }
    .floating-help {
        position: fixed; bottom: 30px; right: 30px; 
        background-color: #2c3e50; color: white; 
        border-radius: 50%; width: 60px; height: 60px; 
        text-align: center; line-height: 60px; font-size: 30px; 
        box-shadow: 2px 2px 10px rgba(0,0,0,0.3); cursor: pointer; z-index: 9999;
        transition: transform 0.2s;
    }
    .floating-help:hover { transform: scale(1.1); background-color: #1a252f; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# SESSION STATE INITIALIZATION
# ==========================================
if 'df' not in st.session_state: st.session_state.df = None
if 'df_clean' not in st.session_state: st.session_state.df_clean = None
if 'data_dict' not in st.session_state: st.session_state.data_dict = {}
if 'report_sections' not in st.session_state: st.session_state.report_sections = []
if 'audit_log' not in st.session_state: st.session_state.audit_log = []

# ==========================================
# UTILITY FUNCTIONS
# ==========================================

def log_action(action):
    """Log user actions for audit trail."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.audit_log.append(f"[{timestamp}] {action}")

def load_data(file):
    """Universal data loader."""
    try:
        ext = file.name.split('.')[-1].lower()
        if ext == 'csv': return pd.read_csv(file)
        elif ext in ['xls', 'xlsx']: return pd.read_excel(file)
        elif ext == 'dta': return pd.read_stata(file)
        elif ext == 'sav': return pd.read_spss(file)
        else: return pd.read_table(file)
    except Exception as e:
        st.error(f"Error loading file: {e}")
        return None

def detect_variable_types(df):
    """Heuristic to detect variable types."""
    types = {}
    for col in df.columns:
        if pd.api.types.is_numeric_dtype(df[col]):
            if df[col].nunique() < 10: types[col] = "Categorical (Ordinal/Binary)"
            else: types[col] = "Continuous"
        elif pd.api.types.is_datetime64_any_dtype(df[col]):
            types[col] = "Date/Time"
        else:
            types[col] = "Categorical (Nominal)"
    return types

def generate_docx_report(title, author, sections):
    """Generates a downloadable DOCX file with embedded images using Kaleido."""
    doc = Document()
    
    # Title Page
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph(f"Author: {author}\nGenerated: {datetime.now().strftime('%Y-%m-%d')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    for section in sections:
        doc.add_heading(section['title'], level=1)
        
        if section['type'] == 'text':
            doc.add_paragraph(section['content'])
            
        elif section['type'] == 'table':
            df = section['content']
            t = doc.add_table(df.shape[0]+1, df.shape[1])
            t.style = 'Table Grid'
            
            # Add header
            for j in range(df.shape[-1]):
                t.cell(0, j).text = str(df.columns[j])
                
            # Add rows
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i+1, j).text = str(df.values[i, j])
            
            doc.add_paragraph("") # Spacing
                    
        elif section['type'] == 'plot':
            try:
                # Use Kaleido to convert Plotly fig to image bytes
                img_bytes = section['content'].to_image(format="png", engine="kaleido", scale=2)
                image_stream = io.BytesIO(img_bytes)
                doc.add_picture(image_stream, width=Inches(6))
                doc.add_paragraph("Figure: " + section['title']).alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Error embedding chart: {str(e)} - Ensure 'kaleido' is installed]")
                
    # Save to IO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def add_to_report(title, content_type, content):
    """Helper to add items to the report buffer."""
    st.session_state.report_sections.append({
        "title": title,
        "type": content_type,
        "content": content
    })
    st.toast(f"Added '{title}' to Report", icon="✅")

# ==========================================
# SIDEBAR
# ==========================================
st.sidebar.title(f"⚕️ {APP_NAME}")
st.sidebar.caption(f"v{VERSION} | Dev: Raja Mashood Elahi")

step = st.sidebar.radio("Workflow", [
    "1. Upload & Setup",
    "2. Data Cleaning",
    "3. Table 1 (Baseline)",
    "4. Exploratory Viz",
    "5. Statistical Tests",
    "6. Regression Models",
    "7. Diagnostic Evaluation",
    "8. Survival Analysis",
    "9. Sample Size (Power)", # NEW
    "10. Report Download"
])

st.sidebar.markdown("---")
st.sidebar.info("💡 **Power Tip:** The report generator now embeds high-res charts directly into Word.")

# ==========================================
# 1. UPLOAD & SETUP
# ==========================================
if step == "1. Upload & Setup":
    st.title("Project Setup")
    
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown("### 📂 Import Data")
        uploaded_file = st.file_uploader("Drop CSV, Excel, SPSS, or Stata files", type=['csv', 'xlsx', 'sav', 'dta'])
        
        if st.button("Load Demo Dataset (Heart Disease)"):
            # Create a more robust dummy dataset
            np.random.seed(42)
            n = 300
            data = {
                'PatientID': range(1, n+1),
                'Age': np.random.normal(60, 10, n).astype(int),
                'Sex': np.random.choice(['Male', 'Female'], n),
                'Cholesterol': np.random.normal(200, 40, n),
                'BMI': np.random.normal(28, 5, n),
                'Treatment': np.random.choice(['Placebo', 'Statin', 'NewDrug'], n),
                'Outcome': np.random.choice(['Recovered', 'No Change'], n, p=[0.6, 0.4]),
                'Disease_Status': np.random.choice([0, 1], n), # 0=No, 1=Yes
                'Test_Result': np.random.normal(0.5, 0.2, n) + (np.random.choice([0, 1], n) * 0.3), # Correlated with disease
                'FollowUp_Days': np.random.exponential(365, n).astype(int),
                'Event_Death': np.random.choice([0, 1], n, p=[0.8, 0.2])
            }
            st.session_state.df = pd.DataFrame(data)
            st.session_state.df_clean = st.session_state.df.copy()
            st.success("Demo data loaded!")

    with col2:
        if st.session_state.df is not None:
            st.markdown("### 🔍 Data Preview")
            st.dataframe(st.session_state.df.head(8), use_container_width=True)
            
            st.markdown("### 🏷️ Variable Types")
            df = st.session_state.df
            detected = detect_variable_types(df)
            
            # Simple Dictionary Editor
            with st.expander("Edit Data Dictionary", expanded=True):
                new_types = {}
                for col in df.columns:
                    c1, c2 = st.columns([1, 2])
                    c1.text(col)
                    options = ["Continuous", "Categorical (Nominal)", "Categorical (Ordinal/Binary)", "Date/Time", "ID/Text"]
                    
                    default_type = detected.get(col, "Categorical (Nominal)")
                    if default_type not in options: default_type = "Categorical (Nominal)"
                    
                    new_types[col] = c2.selectbox(f"Type: {col}", options, 
                                                index=options.index(default_type), 
                                                label_visibility="collapsed")
                st.session_state.data_dict = new_types

# ==========================================
# 2. DATA CLEANING
# ==========================================
elif step == "2. Data Cleaning":
    st.title("Data Cleaning Studio")
    df = st.session_state.df_clean
    if df is None: st.warning("Upload data first."); st.stop()

    t1, t2, t3 = st.tabs(["Missing Data", "Recode/Filter", "Outliers"])

    with t1:
        st.subheader("Imputation")
        miss_cols = df.columns[df.isnull().any()].tolist()
        if miss_cols:
            c1, c2 = st.columns(2)
            target = c1.selectbox("Target Column", miss_cols)
            method = c2.selectbox("Method", ["Mean", "Median", "Mode", "Drop Rows"])
            if st.button("Apply Imputation"):
                if method == "Drop Rows":
                    st.session_state.df_clean = df.dropna(subset=[target])
                else:
                    strategy = 'mean' if method == "Mean" else 'median' if method == "Median" else 'most_frequent'
                    imp = SimpleImputer(strategy=strategy)
                    st.session_state.df_clean[[target]] = imp.fit_transform(df[[target]])
                st.rerun()
        else:
            st.success("No missing values found.")

    with t2:
        st.subheader("Filter Rows")
        f_col = st.selectbox("Filter by", df.columns)
        f_val = st.text_input(f"Keep rows where {f_col} equals (or > for numbers):")
        if st.button("Apply Filter"):
            try:
                # Basic filtering logic
                if pd.api.types.is_numeric_dtype(df[f_col]):
                    st.session_state.df_clean = df[df[f_col] >= float(f_val)]
                else:
                    st.session_state.df_clean = df[df[f_col].astype(str) == f_val]
                st.success(f"Filtered. Rows remaining: {len(st.session_state.df_clean)}")
            except:
                st.error("Filter failed. Check your input.")

    with t3:
        st.subheader("Outlier Removal (IQR Method)")
        num_cols = df.select_dtypes(include=np.number).columns
        target_out = st.selectbox("Check Outliers", num_cols)
        if st.button("Remove Outliers"):
            Q1 = df[target_out].quantile(0.25)
            Q3 = df[target_out].quantile(0.75)
            IQR = Q3 - Q1
            filter_mask = ~((df[target_out] < (Q1 - 1.5 * IQR)) | (df[target_out] > (Q3 + 1.5 * IQR)))
            st.session_state.df_clean = df[filter_mask]
            st.success(f"Removed {len(df) - len(st.session_state.df_clean)} outliers.")

# ==========================================
# 3. TABLE 1
# ==========================================
elif step == "3. Table 1 (Baseline)":
    st.title("Baseline Characteristics (Table 1)")
    df = st.session_state.df_clean
    
    col_strat = st.selectbox("Stratify by (Group)", ["None"] + list(df.columns))
    cols_inc = st.multiselect("Variables", df.columns, default=list(df.columns)[:5])
    
    if st.button("Generate Table"):
        data = []
        for col in cols_inc:
            if col == col_strat: continue
            
            is_numeric = pd.api.types.is_numeric_dtype(df[col]) and df[col].nunique() > 5
            row = {"Characteristic": col}
            
            if col_strat != "None":
                groups = sorted(df[col_strat].dropna().unique())
                p_val = "N/A"
                
                try:
                    group_arrays = [df[df[col_strat] == g][col].dropna() for g in groups]
                    if is_numeric:
                        if len(groups) == 2:
                             _, p = stats.ttest_ind(*group_arrays)
                        else:
                             _, p = stats.f_oneway(*group_arrays)
                        row["P-Value"] = f"{p:.3f}"
                    else:
                        contingency = pd.crosstab(df[col], df[col_strat])
                        _, p, _, _ = stats.chi2_contingency(contingency)
                        row["P-Value"] = f"{p:.3f}"
                except:
                    row["P-Value"] = "-"

                for g in groups:
                    sub = df[df[col_strat] == g]
                    if is_numeric:
                        row[str(g)] = f"{sub[col].mean():.1f} ± {sub[col].std():.1f}"
                    else:
                        top = sub[col].mode()[0] if not sub[col].empty else "N/A"
                        count = sub[col].value_counts().max()
                        pct = (count/len(sub))*100
                        row[str(g)] = f"{count} ({pct:.1f}%)"
            else:
                if is_numeric:
                    row["Total"] = f"{df[col].mean():.1f} ± {df[col].std():.1f}"
                else:
                    row["Total"] = f"n={len(df)}"

            data.append(row)
        
        res_df = pd.DataFrame(data)
        st.dataframe(res_df, use_container_width=True)
        add_to_report("Table 1: Baseline Characteristics", "table", res_df)

# ==========================================
# 4. EXPLORATORY VIZ
# ==========================================
elif step == "4. Exploratory Viz":
    st.title("Exploratory Data Analysis")
    df = st.session_state.df_clean
    
    viz_type = st.selectbox("Chart Type", ["Scatter Plot", "Box Plot", "Violin Plot", "Histogram", "Heatmap", "Bar Chart"])
    
    c1, c2, c3 = st.columns(3)
    x_var = c1.selectbox("X Axis", df.columns)
    y_var = c2.selectbox("Y Axis (Optional)", ["None"] + list(df.columns))
    color_var = c3.selectbox("Color By", ["None"] + list(df.columns))
    
    fig = None
    y = None if y_var == "None" else y_var
    color = None if color_var == "None" else color_var
    
    if viz_type == "Scatter Plot" and y:
        fig = px.scatter(df, x=x_var, y=y, color=color, trendline="ols", template="plotly_white")
    elif viz_type == "Box Plot":
        fig = px.box(df, x=x_var, y=y, color=color, template="plotly_white")
    elif viz_type == "Violin Plot":
        fig = px.violin(df, x=x_var, y=y, color=color, box=True, template="plotly_white")
    elif viz_type == "Histogram":
        fig = px.histogram(df, x=x_var, color=color, marginal="box", template="plotly_white")
    elif viz_type == "Bar Chart":
        fig = px.bar(df, x=x_var, y=y, color=color, template="plotly_white")
    elif viz_type == "Heatmap":
        numeric_df = df.select_dtypes(include=np.number)
        fig = px.imshow(numeric_df.corr(), text_auto=True, template="plotly_white", aspect="auto")
        
    if fig:
        st.plotly_chart(fig, use_container_width=True)
        if st.button("Add Plot to Report"):
            add_to_report(f"{viz_type}: {x_var} vs {y_var}", "plot", fig)

# ==========================================
# 5. STATISTICAL TESTS (POWERHOUSE)
# ==========================================
elif step == "5. Statistical Tests":
    st.title("Smart Statistical Engine (Assumption Aware)")
    df = st.session_state.df_clean
    
    c1, c2 = st.columns(2)
    v1 = c1.selectbox("Variable A (Outcome)", df.columns)
    v2 = c2.selectbox("Variable B (Group)", df.columns)
    
    if st.button("Run Auto-Test"):
        st.markdown("### 📊 Analysis Results")
        
        is_num_1 = pd.api.types.is_numeric_dtype(df[v1])
        is_num_2 = pd.api.types.is_numeric_dtype(df[v2])
        
        # Scenario 1: Continuous Outcome, Categorical Group
        if is_num_1 and not is_num_2:
            groups = df[v2].unique()
            group_data = [df[df[v2]==g][v1].dropna() for g in groups]
            
            # Assumption Checking
            st.subheader("1. Assumption Checks")
            
            # Normality
            normality_p = stats.shapiro(df[v1].dropna().sample(min(50, len(df))))[1]
            st.write(f"- **Shapiro-Wilk Normality:** p={normality_p:.4f} " + 
                     ("✅ (Normal)" if normality_p > 0.05 else "⚠️ (Non-Normal)"))
            
            # Homogeneity
            try:
                levene_p = stats.levene(*group_data)[1]
                st.write(f"- **Levene's Variance:** p={levene_p:.4f} " + 
                         ("✅ (Equal)" if levene_p > 0.05 else "⚠️ (Unequal)"))
            except: levene_p = 1.0

            st.markdown("---")
            st.subheader("2. Test Result")
            
            if len(groups) == 2:
                if normality_p > 0.05 and levene_p > 0.05:
                    test_name = "Student's T-Test"
                    stat, p = stats.ttest_ind(*group_data)
                else:
                    test_name = "Mann-Whitney U Test (Non-Parametric)"
                    stat, p = stats.mannwhitneyu(*group_data)
                
                st.metric(f"{test_name} P-Value", f"{p:.4f}", delta="Significant" if p<0.05 else "NS")
                
            else: # > 2 Groups
                if normality_p > 0.05:
                    test_name = "One-Way ANOVA"
                    stat, p = stats.f_oneway(*group_data)
                else:
                    test_name = "Kruskal-Wallis H Test"
                    stat, p = stats.kruskal(*group_data)
                
                st.metric(f"{test_name} P-Value", f"{p:.4f}", delta="Significant" if p<0.05 else "NS")
                
                if p < 0.05 and test_name == "One-Way ANOVA":
                    st.write("**Post-Hoc Analysis (Tukey's HSD):**")
                    tukey = pairwise_tukeyhsd(endog=df[v1].dropna(), groups=df[v2].dropna(), alpha=0.05)
                    st.text(tukey)

            fig = px.box(df, x=v2, y=v1, points="all", title=f"{test_name} Result")
            st.plotly_chart(fig)
            add_to_report(f"Comparison: {v1} by {v2}", "plot", fig)
            
        # Scenario 2: Categorical vs Categorical
        elif not is_num_1 and not is_num_2:
            st.info("Categorical vs Categorical → Chi-Square Test")
            ct = pd.crosstab(df[v1], df[v2])
            chi2, p, dof, expected = stats.chi2_contingency(ct)
            st.metric("Chi-Square P-Value", f"{p:.4f}")
            
            fig = px.imshow(ct, text_auto=True, title="Contingency Heatmap")
            st.plotly_chart(fig)
            add_to_report(f"Chi-Square: {v1} vs {v2}", "plot", fig)
            
        # Scenario 3: Correlation
        elif is_num_1 and is_num_2:
            r, p = stats.pearsonr(df[v1].dropna(), df[v2].dropna())
            st.metric(f"Pearson r", f"{r:.3f}")
            st.metric(f"P-Value", f"{p:.4f}")
            fig = px.scatter(df, x=v1, y=v2, trendline="ols")
            st.plotly_chart(fig)
            add_to_report(f"Correlation: {v1} vs {v2}", "plot", fig)

# ==========================================
# 6. REGRESSION
# ==========================================
elif step == "6. Regression Models":
    st.title("Regression Modeling")
    df = st.session_state.df_clean
    
    model_type = st.selectbox("Model", ["Linear (OLS)", "Logistic (Binary)"])
    target = st.selectbox("Outcome Variable (Y)", df.columns)
    preds = st.multiselect("Predictors (X)", [c for c in df.columns if c != target])
    
    if st.button("Fit Model") and preds:
        formula = f"{target} ~ {' + '.join(preds)}"
        try:
            if model_type == "Linear (OLS)":
                model = smf.ols(formula, data=df).fit()
            else:
                model = smf.logit(formula, data=df).fit()
            
            st.text(model.summary().as_text())
            
            # Forest Plot Data
            res_df = pd.DataFrame({
                "Coef": model.params,
                "Lower": model.conf_int()[0],
                "Upper": model.conf_int()[1]
            })
            if "Logistic" in model_type: res_df = np.exp(res_df) # Odds Ratios
            
            # Forest Plot
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=res_df['Coef'], y=res_df.index,
                error_x=dict(type='data', array=res_df['Upper']-res_df['Coef'], arrayminus=res_df['Coef']-res_df['Lower']),
                mode='markers', marker=dict(color='black', size=10)
            ))
            fig.update_layout(title=f"Forest Plot ({'Odds Ratios' if 'Logistic' in model_type else 'Coefficients'})", template="plotly_white")
            line_val = 1 if "Logistic" in model_type else 0
            fig.add_vline(x=line_val, line_dash="dash", line_color="red")
            
            st.plotly_chart(fig)
            add_to_report(f"Regression Model ({model_type})", "plot", fig)
            
        except Exception as e:
            st.error(f"Modeling Error: {e}")

# ==========================================
# 7. DIAGNOSTIC EVALUATION
# ==========================================
elif step == "7. Diagnostic Evaluation":
    st.title("Diagnostic Test Evaluation")
    df = st.session_state.df_clean
    
    c1, c2 = st.columns(2)
    gold_standard = c1.selectbox("Gold Standard (Binary)", df.columns)
    test_var = c2.selectbox("Test Variable", df.columns)
    
    if st.button("Run Diagnostic Analysis"):
        if df[gold_standard].nunique() != 2:
            st.error("Gold Standard must be binary.")
        else:
            y_true = df[gold_standard]
            y_score = df[test_var]
            
            if pd.api.types.is_numeric_dtype(y_score) and y_score.nunique() > 2:
                # ROC Curve
                fpr, tpr, thresholds = roc_curve(y_true, y_score, pos_label=1)
                roc_auc = auc(fpr, tpr)
                
                fig = px.area(x=fpr, y=tpr, title=f'ROC Curve (AUC = {roc_auc:.3f})',
                              labels=dict(x='1 - Specificity', y='Sensitivity'))
                fig.add_shape(type='line', line=dict(dash='dash'), x0=0, x1=1, y0=0, y1=1)
                st.plotly_chart(fig)
                add_to_report(f"ROC Curve: {test_var}", "plot", fig)
                
                # Youden's J
                j_scores = tpr - fpr
                optimal_idx = np.argmax(j_scores)
                st.metric("Optimal Cutoff", f"{thresholds[optimal_idx]:.2f}")
                
            else:
                # Confusion Matrix
                cm = confusion_matrix(y_true, y_score)
                tn, fp, fn, tp = cm.ravel()
                
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Sensitivity", f"{tp/(tp+fn):.2%}")
                m2.metric("Specificity", f"{tn/(tn+fp):.2%}")
                m3.metric("PPV", f"{tp/(tp+fp):.2%}")
                m4.metric("NPV", f"{tn/(tn+fn):.2%}")
                
                z = [[tn, fp], [fn, tp]]
                fig_cm = ff.create_annotated_heatmap(z, x=['Neg', 'Pos'], y=['No Disease', 'Disease'], colorscale='Blues')
                st.plotly_chart(fig_cm)
                add_to_report(f"Confusion Matrix: {test_var}", "plot", fig_cm)

# ==========================================
# 8. SURVIVAL
# ==========================================
elif step == "8. Survival Analysis":
    st.title("Survival Analysis")
    df = st.session_state.df_clean
    
    c1, c2, c3 = st.columns(3)
    time_col = c1.selectbox("Time Column", df.columns)
    event_col = c2.selectbox("Event Column (0/1)", df.columns)
    group_col = c3.selectbox("Compare Groups", ["None"] + list(df.columns))
    
    if st.button("Plot Survival"):
        kmf = KaplanMeierFitter()
        fig = go.Figure()
        
        if group_col == "None":
            kmf.fit(df[time_col], df[event_col], label="Overall")
            fig.add_trace(go.Scatter(x=kmf.timeline, y=kmf.survival_function_.iloc[:,0], mode='lines', name="Overall"))
        else:
            for name, grouped_df in df.groupby(group_col):
                kmf.fit(grouped_df[time_col], grouped_df[event_col], label=str(name))
                fig.add_trace(go.Scatter(x=kmf.timeline, y=kmf.survival_function_.iloc[:,0], mode='lines', name=str(name)))
                
        fig.update_layout(title="Kaplan-Meier Curve", xaxis_title="Time", yaxis_title="Survival Probability", template="plotly_white")
        st.plotly_chart(fig)
        add_to_report("Kaplan-Meier Curve", "plot", fig)

# ==========================================
# 9. SAMPLE SIZE CALCULATOR (NEW)
# ==========================================
elif step == "9. Sample Size (Power)":
    st.title("Sample Size Calculator")
    st.info("Calculate sample size for comparing two means (Independent T-test).")
    
    c1, c2, c3 = st.columns(3)
    effect_size = c1.number_input("Effect Size (Cohen's d)", 0.2, 2.0, 0.5, step=0.1)
    power = c2.number_input("Power (1 - beta)", 0.5, 0.99, 0.8, step=0.05)
    alpha = c3.number_input("Alpha (Sig. Level)", 0.01, 0.1, 0.05, step=0.01)
    
    if st.button("Calculate N"):
        analysis = TTestIndPower()
        result = analysis.solve_power(effect_size=effect_size, power=power, alpha=alpha, ratio=1.0)
        st.metric("Required Sample Size (Per Group)", f"{int(np.ceil(result))}")
        st.caption("Total sample size ≈ " + str(int(np.ceil(result)) * 2))

# ==========================================
# 10. REPORT DOWNLOAD
# ==========================================
elif step == "10. Report Download":
    st.title("📄 Generate Final Report")
    
    st.write(f"**Items in Report:** {len(st.session_state.report_sections)}")
    
    col1, col2 = st.columns(2)
    title = col1.text_input("Title", "Analysis Report")
    author = col2.text_input("Author", "Investigator")
    
    c_html, c_docx = st.columns(2)
    
    with c_html:
        st.markdown("#### Interactive HTML")
        if st.button("Download HTML"):
            # HTML generation (Simplified for brevity)
            html_content = f"<h1>{title}</h1><hr>"
            for section in st.session_state.report_sections:
                html_content += f"<h3>{section['title']}</h3>"
                if section['type'] == 'plot':
                    html_content += section['content'].to_html(full_html=False, include_plotlyjs='cdn')
                elif section['type'] == 'table':
                    html_content += section['content'].to_html()
            
            b64 = base64.b64encode(html_content.encode()).decode()
            href = f'<a href="data:text/html;base64,{b64}" download="report.html">Download HTML</a>'
            st.markdown(href, unsafe_allow_html=True)

    with c_docx:
        st.markdown("#### Word Document (with Images)")
        if st.button("Download DOCX"):
            try:
                docx_file = generate_docx_report(title, author, st.session_state.report_sections)
                st.download_button(
                    label="📥 Download DOCX",
                    data=docx_file,
                    file_name="MedStat_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Generation Error: {e}")

# ==========================================
# FLOATING HELP
# ==========================================
st.markdown('<div class="floating-help">?</div>', unsafe_allow_html=True)